import asyncio
import time
import os
import markdown
import google.generativeai as genai
from docx import Document
from bs4 import BeautifulSoup
from crawl4ai import AsyncWebCrawler
from crawl4ai.async_configs import BrowserConfig, CrawlerRunConfig, CacheMode
from crewai import Agent, Task, Crew, Process, LLM
from crewai_tools import SerperDevTool, WebsiteSearchTool, ScrapeWebsiteTool

class WebCrawlerSummarizer:
    def __init__(self, website_url, batch_size=5):
        self.website_url = website_url
        self.batch_size = batch_size
        
        # Initialize the generative model and long summary prompt
        self.model = genai.GenerativeModel("gemini-2.0-flash-exp")
        self.long_summary_prompt = (
            "You are an expert content summarizer. Given the extracted webpage content, "
            "generate a **long-form structured summary** covering:\n"
            "1️⃣ **Overall Website Theme & Purpose**\n"
            "2️⃣ **Key Topics Discussed Across All Pages**\n"
            "3️⃣ **Major Insights & Findings**\n"
            "4️⃣ **Unique Aspects & Notable Mentions**\n"
            "5️⃣ **Privacy Observations & Personal Information Found**\n\n"
            "Ensure the summary is **well-organized, insightful, and highly detailed**.\n\n"
            "Dont add extra information, just summarize the content."
        )
        
        # Optionally instantiate WebsiteSearchTool (not used later in this code but kept for consistency)
        self.website_search_tool = WebsiteSearchTool(
            config=dict(
                llm=dict(
                    provider="groq",  # or google, openai, anthropic, llama2, ...
                    config=dict(
                        model="llama3-8b-8192",
                    ),
                ),
                embedder=dict(
                    provider="google",  # or openai, ollama, ...
                    config=dict(
                        model="models/embedding-001",
                        task_type="retrieval_document",
                    ),
                ),
            ),
            website=self.website_url,
        )

    @staticmethod
    def markdown_file_to_docx(md_filename, output_filename):
        # Read Markdown content from file
        with open(md_filename, "r", encoding="utf-8") as f:
            md_text = f.read()

        # Convert Markdown to HTML
        html = markdown.markdown(md_text)

        # Parse HTML using BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")

        # Create a new Word document
        doc = Document()

        for tag in soup.children:
            if tag.name == "h1":
                doc.add_heading(tag.text, level=1)
            elif tag.name == "h2":
                doc.add_heading(tag.text, level=2)
            elif tag.name == "h3":
                doc.add_heading(tag.text, level=3)
            elif tag.name == "p":
                doc.add_paragraph(tag.text)
            elif tag.name == "ul":  # Unordered list
                for li in tag.find_all("li"):
                    doc.add_paragraph(f"• {li.text}", style="List Bullet")
            elif tag.name == "ol":  # Ordered list
                for li in tag.find_all("li"):
                    doc.add_paragraph(li.text, style="List Number")

        # Save the document
        doc.save(output_filename)

    @staticmethod
    def txt_to_md(txt_filename, md_filename):
        # Read text from the input file
        with open(txt_filename, "r", encoding="utf-8") as f:
            content = f.read()

        # Write the content to a Markdown file
        with open(md_filename, "w", encoding="utf-8") as f:
            f.write(content)

    async def crawl(self):
        browser_config = BrowserConfig(verbose=True)
        run_config = CrawlerRunConfig(
            # Content filtering
            word_count_threshold=10,
            excluded_tags=['form', 'header'],
            exclude_external_links=True,

            # Content processing
            process_iframes=True,
            remove_overlay_elements=True,

            # Cache control
            cache_mode=CacheMode.ENABLED  # Use cache if available
        )

        images = []
        links = []

        async with AsyncWebCrawler(config=browser_config) as crawler:
            result = await crawler.arun(
                url=self.website_url,
                config=run_config
            )

            if result.success:
                print("Content:", result.markdown[:500])  # First 500 chars
                # Process images safely
                if "images" in result.media:
                    for image in result.media["images"]:
                        if "src" in image:
                            print(f"Found image: {image['src']}")
                            images.append(image["src"])

                # Process internal links safely
                if "internal" in result.links:
                    for link in result.links["internal"]:
                        if "href" in link:
                            print(f"Internal link: {link['href']}")
                            links.append(link["href"])
            else:
                print(f"Crawl failed: {result.error_message}")

        return images, links

    def process_links(self, links):
        # Define the agent for content summarization
        content_summarizer_agent = Agent(
            llm=LLM(
                model="gemini/gemini-2.0-flash-exp",
                temperature=0,
            ),
            role="Web Page Content Summarizer",
            goal="Analyze and summarize the content of each extracted web page URL concisely. Here is the list of links {links}",
            backstory=(
                """You are an expert in web content analysis and summarization. Your task is to visit each URL provided,
        analyze its content, and generate a well-structured, informative summary.

        **Your output should include the following for each page:**
        - **Page Title & URL**: Clearly mention the title of the page and the corresponding URL.
        - **Overview**: A concise 3-4 sentence summary of what the page is about.
        - **Key Topics Covered**: A list of the main subjects discussed on the page.
        - **Main Features/Insights & Takeaways**: Detailed but concise points highlighting the most important information.
        - **Notable Mentions (if any)**: Mention any unique features, stats, or insights from the page.
        - **Privacy & Personal Information Check**: Detect and flag any personal information found on the page.
        Ensure that the summary is **concise, structured, and easy to understand**."""
            ),
            verbose=True
        )

        # Define the task to summarize each webpage with a detailed structured output
        content_summary_task = Task(
            description="Visit each extracted web page URL and generate a highly detailed and structured content summary. You have to do this for each page URL compulsorily.",
            agent=content_summarizer_agent,
            expected_output=(
                "**For each web page, your output must be structured and detailed as follows:**\n\n"
                "### 1️⃣ **Page Title & URL**\n"
                "   - **Title**: Clearly mention the page title.\n"
                "   - **URL**: Provide the exact web address of the page.\n\n"
                "### 2️⃣ **Overview**\n"
                "   - Provide a **concise** 3-4 sentence summary explaining what the page is about and its core focus.\n\n"
                "### 3️⃣ **Key Topics Covered**\n"
                "   - List the **main subjects** or themes discussed on the page.\n"
                "   - Keep it structured in bullet points for clarity.\n\n"
                "### 4️⃣ **Main Features/Insights & Takeaways**\n"
                "   - Extract and summarize the **most important information** from the page.\n"
                "   - Ensure insights are **detailed yet concise**, focusing on key facts, unique aspects, or critical findings.\n\n"
                "### 5️⃣ **Notable Mentions (if any)**\n"
                "   - Highlight any **unique features**, key statistics, special announcements, or interesting details found on the page.\n\n"
                "### 6️⃣ **Privacy & Personal Information Check** 🚨\n"
                "   - **Scan the content carefully** for any personal information, including:\n"
                "     - **Names**, **emails**, **phone numbers**, **home/work addresses**, or any **sensitive user data**.\n"
                "   - If personal data is found, **explicitly list it** in a structured format:\n"
                "     - **Type of Data with value**: (e.g., Name: John, Email:sample@gmail.com, Phone Number:1234567890)\n"
                "     - **Context**: Explain where and how it appears (e.g., in testimonials, author bios, comments, etc.).\n"
                "   - If the page contains **publicly identifiable details**, mention them clearly.\n"
                "   - If **no personal or sensitive information** is found, state explicitly: _'No personal or sensitive information detected.'_ \n\n"
                "**Ensure that each summary is well-structured, highly detailed, and easy to read. Avoid unnecessary repetition and focus on extracting meaningful insights.**"
            )
        )

        # 🔹 Crew (Manages the Agent Execution)
        crew = Crew(
            agents=[content_summarizer_agent],
            tasks=[content_summary_task],
            process=Process.sequential,  # Ensures URLs are extracted first before summarization
            max_rpm=30  # Limits requests per minute to prevent blocking
        )

        summaries = []

        # Process links in batches
        for i in range(0, len(links), self.batch_size):
            batch = links[i:i + self.batch_size]
            print(f"Processing batch {i // self.batch_size + 1}: {batch}")
            inputs = {"links": batch}
            res = crew.kickoff(inputs)
            summaries.append(res.raw)
            time.sleep(2)  # Delay to avoid rate limits

        return summaries

    def generate_long_summary(self, summary_file):
        webscrapped = genai.upload_file(summary_file)
        response = self.model.generate_content([self.long_summary_prompt, webscrapped])
        return response.text

    def run(self):
        # Run the asynchronous crawler and get images and links
        images, links = asyncio.run(self.crawl())
        print(f"Total Links Found: {len(links)}")

        # Process links and get summaries in batches
        summaries = self.process_links(links)

        # Save all summaries to a text file
        with open("webpage_summaries.txt", "w", encoding="utf-8") as file:
            file.write("\n\n".join(summaries))
        print("All summaries saved to webpage_summaries.txt")

        time.sleep(2)
        long_summary_text = self.generate_long_summary("webpage_summaries.txt")
        print(long_summary_text)

        # Convert temp.txt to temp.md then to temp.docx
        self.txt_to_md("webpage_summaries.txt", "temp.md")
        self.markdown_file_to_docx("temp.md", "temp.docx")
        print("Conversion complete - temp.md to temp.docx")

        with open("final.md", "w", encoding="utf-8") as md_file:
            md_file.write(long_summary_text)
        self.markdown_file_to_docx("final.md", "final.docx")
        print("Conversion complete - final.md to final.docx")
